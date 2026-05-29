"""
Output Writer Agent - Handles saving and formatting output documents.

Renders the Markdown produced by the SecretaryAgent into a clean, professional
Word document suitable for status reporting. Markdown is parsed into native Word
constructs (heading styles, bullet lists, inline bold/italic, shaded tables) so
the output is presentation-ready with no manual clean-up.
"""

import re
from typing import List, Dict, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.base_agent import BaseAgent


# Corporate palette tuned for status reporting.
ACCENT = RGBColor(0x1F, 0x49, 0x7D)        # deep blue - titles & section headings
SUBHEAD = RGBColor(0x2E, 0x74, 0xB5)       # mid blue - sub headings
MUTED = RGBColor(0x60, 0x60, 0x60)         # grey - metadata
HEADER_FILL = "1F497D"                      # table header shading (hex, no #)
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)    # table header text
BAND_FILL = "EAF1F8"                        # alternating row shading

BODY_FONT = "Calibri"


class OutputWriterAgent(BaseAgent):
    """Agent responsible for saving and formatting output documents"""

    def __init__(self):
        super().__init__("OutputWriter")

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #
    def execute(self, results: List[Dict[str, str]], output_file: str = "Meeting_Minutes_Summary.docx") -> bool:
        """
        Save processed results to a Word document.

        Args:
            results: List of dictionaries with 'file_name' and 'summary' keys
            output_file: Name of the output file

        Returns:
            True if successful, False otherwise
        """
        try:
            self.log(f"Writing results to {output_file}...")

            doc = Document()
            self._configure_styles(doc)
            self._add_cover(doc, len(results))

            for i, result in enumerate(results, 1):
                if i > 1:
                    doc.add_page_break()
                self._add_section_title(doc, result["file_name"])
                self._render_markdown(doc, result["summary"])

            doc.save(output_file)
            self.log(f"Successfully saved results to {output_file}")
            return True

        except Exception as e:
            self.log(f"Error writing output: {e}", "ERROR")
            return False

    # ------------------------------------------------------------------ #
    # Document scaffolding
    # ------------------------------------------------------------------ #
    def _configure_styles(self, doc: Document) -> None:
        """Set a clean base font and consistent heading styling."""
        normal = doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(4)
        normal.paragraph_format.line_spacing = 1.12

        heading_specs = {
            "Heading 1": (16, ACCENT, True),
            "Heading 2": (13, SUBHEAD, True),
            "Heading 3": (11.5, SUBHEAD, True),
        }
        for name, (size, color, bold) in heading_specs.items():
            try:
                style = doc.styles[name]
            except KeyError:
                continue
            style.font.name = BODY_FONT
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.font.bold = bold
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(4)

    def _add_cover(self, doc: Document, count: int) -> None:
        """Title block at the top of the report."""
        title = doc.add_paragraph()
        run = title.add_run("Meeting Minutes Summary")
        run.font.name = BODY_FONT
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        title.paragraph_format.space_after = Pt(2)

        meta = doc.add_paragraph()
        label = "transcript" if count == 1 else "transcripts"
        run = meta.add_run(
            f"Generated {datetime.now().strftime('%d %b %Y, %H:%M')}  ·  {count} {label}"
        )
        run.font.size = Pt(9.5)
        run.font.color.rgb = MUTED
        meta.paragraph_format.space_after = Pt(2)

        self._add_rule(doc, ACCENT, size=12)

    def _add_section_title(self, doc: Document, file_name: str) -> None:
        """Per-transcript header derived from the source file name."""
        name = re.sub(r"\.docx$", "", file_name, flags=re.IGNORECASE)
        heading = doc.add_heading(name, level=1)
        heading.paragraph_format.space_before = Pt(0)

    # ------------------------------------------------------------------ #
    # Markdown rendering
    # ------------------------------------------------------------------ #
    def _render_markdown(self, doc: Document, text: str) -> None:
        """Convert the Markdown summary into native Word constructs."""
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            raw = lines[i]
            stripped = raw.strip()

            # Skip blank lines and stray document-level # MEETING MINUTES title.
            if not stripped:
                i += 1
                continue

            # Markdown table: a run of pipe-delimited lines.
            if self._is_table_row(stripped):
                block = []
                while i < len(lines) and self._is_table_row(lines[i].strip()):
                    block.append(lines[i].strip())
                    i += 1
                self._add_table(doc, block)
                continue

            # Horizontal rule.
            if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
                self._add_rule(doc, RGBColor(0xCC, 0xCC, 0xCC), size=6)
                i += 1
                continue

            # Headings (#, ##, ###, ...).
            heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if heading_match:
                self._add_heading(doc, heading_match)
                i += 1
                continue

            # Bullet list items (-, *, +) with indentation -> nesting level.
            bullet_match = re.match(r"^(\s*)[-*+]\s+(.*)$", raw)
            if bullet_match:
                indent = len(bullet_match.group(1).replace("\t", "  "))
                level = min(indent // 2, 2)
                self._add_bullet(doc, bullet_match.group(2), level)
                i += 1
                continue

            # Numbered list items.
            ordered_match = re.match(r"^(\s*)\d+[.)]\s+(.*)$", raw)
            if ordered_match:
                para = doc.add_paragraph(style="List Number")
                self._add_inline(para, ordered_match.group(2))
                i += 1
                continue

            # Plain paragraph.
            para = doc.add_paragraph()
            self._add_inline(para, stripped)
            i += 1

    def _add_heading(self, doc: Document, match: re.Match) -> None:
        hashes, content = match.group(1), self._strip_inline_markers(match.group(2))
        level = len(hashes)
        # The model emits a top-level "# MEETING MINUTES" banner; the cover already
        # provides the report title, so render it as a compact subheading instead.
        if level == 1:
            level = 2
        doc.add_heading(content, level=min(level, 3))

    def _add_bullet(self, doc: Document, content: str, level: int) -> None:
        style = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
        try:
            para = doc.add_paragraph(style=style)
        except KeyError:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        self._add_inline(para, content)

    # ------------------------------------------------------------------ #
    # Inline (bold / italic / code) formatting
    # ------------------------------------------------------------------ #
    _INLINE_RE = re.compile(
        r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`)"
    )

    def _add_inline(self, paragraph, text: str) -> None:
        """Add text to a paragraph, honouring **bold**, *italic*, and `code`."""
        for token in self._INLINE_RE.split(text):
            if not token:
                continue
            if (token.startswith("**") and token.endswith("**")) or (
                token.startswith("__") and token.endswith("__")
            ):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif (token.startswith("*") and token.endswith("*")) or (
                token.startswith("_") and token.endswith("_")
            ):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            elif token.startswith("`") and token.endswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = "Consolas"
            else:
                paragraph.add_run(token)

    @staticmethod
    def _strip_inline_markers(text: str) -> str:
        return re.sub(r"[*_`]+", "", text).strip()

    # ------------------------------------------------------------------ #
    # Tables
    # ------------------------------------------------------------------ #
    @staticmethod
    def _is_table_row(line: str) -> bool:
        return line.startswith("|") and line.count("|") >= 2

    @staticmethod
    def _split_row(line: str) -> List[str]:
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    @staticmethod
    def _is_separator_row(cells: List[str]) -> bool:
        return all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)

    def _add_table(self, doc: Document, block: List[str]) -> None:
        rows = [self._split_row(r) for r in block]
        # Drop the Markdown header/body separator row (|---|---|).
        rows = [r for r in rows if not self._is_separator_row(r)]
        if not rows:
            return

        n_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=0, cols=n_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for r_idx, cells in enumerate(rows):
            cells = cells + [""] * (n_cols - len(cells))
            row = table.add_row()
            is_header = r_idx == 0
            for c_idx, cell_text in enumerate(cells):
                cell = row.cells[c_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.space_before = Pt(2)
                self._add_inline(para, self._strip_inline_markers(cell_text)
                                 if is_header else cell_text)
                if is_header:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = HEADER_TEXT
                    self._shade_cell(cell, HEADER_FILL)
                elif r_idx % 2 == 0:
                    self._shade_cell(cell, BAND_FILL)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    @staticmethod
    def _shade_cell(cell, hex_fill: str) -> None:
        """Apply background shading to a table cell."""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        tc_pr.append(shd)

    # ------------------------------------------------------------------ #
    # Misc helpers
    # ------------------------------------------------------------------ #
    @staticmethod
    def _add_rule(doc: Document, color: RGBColor, size: int = 6) -> None:
        """Add a horizontal divider line below an empty paragraph."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(6)
        p_pr = para._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
